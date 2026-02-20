from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document


LIMIT_LINE_PATTERN = re.compile(
    r"^\s*(?P<name>[A-Za-z0-9\-\(\)\/\s]+?)\s+(?P<pressure>-?\d+(?:\.\d+)?)\s+(?P<temperature>-?\d+(?:\.\d+)?(?:\s*to\s*-?\d+(?:\.\d+)?)?)\s*$",
    flags=re.IGNORECASE,
)


def extract_sop_text(docx_path: str | Path) -> str:
    docx_path = Path(docx_path)
    doc = Document(docx_path)
    lines: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def extract_design_limits(docx_path: str | Path) -> list[dict[str, Any]]:
    docx_path = Path(docx_path)
    doc = Document(docx_path)

    extracted: list[dict[str, Any]] = []
    seen = set()

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if len(cells) < 3:
                continue
            name, pressure, temperature = cells[0], cells[1], cells[2]
            if not _looks_like_limit_row(name, pressure, temperature):
                continue
            key = name.upper()
            if key in seen:
                continue
            extracted.append(
                {
                    "name": name,
                    "pressure_psig": pressure,
                    "temperature_f": temperature,
                    "source": "table",
                }
            )
            seen.add(key)

    text_lines = extract_sop_text(docx_path).splitlines()
    for line in text_lines:
        match = LIMIT_LINE_PATTERN.match(line)
        if not match:
            continue
        name = match.group("name").strip()
        key = name.upper()
        if key in seen:
            continue
        extracted.append(
            {
                "name": name,
                "pressure_psig": match.group("pressure").strip(),
                "temperature_f": match.group("temperature").strip(),
                "source": "regex_line",
            }
        )
        seen.add(key)

    return extracted


def _looks_like_limit_row(name: str, pressure: str, temperature: str) -> bool:
    if not name or not pressure or not temperature:
        return False
    name_lower = name.lower()
    if "pressure" in name_lower or "temperature" in name_lower:
        return False
    has_pressure_num = any(ch.isdigit() for ch in pressure)
    has_temp_num = any(ch.isdigit() for ch in temperature) or "-" in temperature
    return has_pressure_num and has_temp_num
